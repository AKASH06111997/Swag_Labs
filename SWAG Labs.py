import os
import time
from datetime import datetime
import docx
from docx.shared import Inches
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

def initialize_driver():
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()))
    driver.maximize_window()
    return driver

def enter_text_slowly(element, text):
    for char in text:
        element.send_keys(char)
        time.sleep(0.1)
def pass_screenshot(doc, driver, idx, message, doc_name):
    screenshot_dir = r"C:\Users\91951\AppData\Local\Programs\Python\Python312\screenshots"
    os.makedirs(screenshot_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    screenshot_path = os.path.join(screenshot_dir, f"screenshot_{timestamp}.png")
    driver.save_screenshot(screenshot_path)
    para = doc.add_paragraph()
    run_id = para.add_run(f"Test Case {idx}: ")
    run_message = para.add_run(message)
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_picture(screenshot_path, width=Inches(7))
    doc.save(doc_name)
    os.remove(screenshot_path)
class Login:
    @staticmethod
    def login_url(driver, username, password):
        driver.get("https://www.saucedemo.com/")

        WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, "//*[@id='user-name']")))
        username_field = driver.find_element(By.XPATH, "//*[@id='user-name']")
        enter_text_slowly(username_field, username)

        password_field = driver.find_element(By.XPATH, "//*[@id='password']")
        enter_text_slowly(password_field, password)

        login_button = driver.find_element(By.XPATH, "//*[@id='login-button']")
        login_button.click()
        time.sleep(5)
def successful_login(driver):
    Login.login_url(driver, "standard_user", "secret_sauce")
    time.sleep(10)
    try:
        status_message = WebDriverWait(driver, 5).until(
            EC.visibility_of_element_located((By.XPATH, "//*[@id='header_container']/div[1]/div[2]/div"))
        )
        if "Swag Labs" in status_message.text:
            message = "Successful login test passed"
        else:
            message = f"Unexpected message: {status_message.text}"
    except Exception as e:
        message = f"Login failed or took too long to load: {e}"
    return message
def invalid_username(driver):
    time.sleep(5)
    logmenu_button = WebDriverWait(driver, 10).until(
        EC.element_to_be_clickable((By.XPATH, "//*[@id='react-burger-menu-btn']"))
    )
    driver.execute_script("arguments[0].scrollIntoView(true);", logmenu_button)
    logmenu_button.click()
    time.sleep(5)

    # Click the logout button
    logout_button = WebDriverWait(driver, 10).until(
        EC.element_to_be_clickable((By.XPATH, "//*[@id='logout_sidebar_link']"))
    )
    driver.execute_script("arguments[0].scrollIntoView(true);", logout_button)
    logout_button.click()
    time.sleep(5)

    Login.login_url(driver, "locked_out_user", "secret_sauce")

    try:
        status_message = WebDriverWait(driver, 5).until(
            EC.visibility_of_element_located((By.XPATH, "//*[@id='login_button_container']/div/form/div[3]/h3"))
        )
        if "Epic sadface: Sorry, this user has been locked out." in status_message.text:
            message = "Sorry, this user has been Locked"
        else:
            message = f"Unexpected message: {status_message.text}"
    except Exception as e:
        message = f"Error occurred while checking success message: {e}"
    return message
def Order_a_product(driver):
    Login.login_url(driver, "standard_user", "secret_sauce")

    filter_button = WebDriverWait(driver, 10).until(
        EC.element_to_be_clickable((By.XPATH, "//*[@id='header_container']/div[2]/div/span/select"))
    )
    filter_button.click()
    time.sleep(2)

    filter_dropdown_option = WebDriverWait(driver, 10).until(
        EC.element_to_be_clickable((By.XPATH, "//*[@id='header_container']/div[2]/div/span/select/option[4]"))
    )
    filter_dropdown_option.click()
    time.sleep(2)

    add_product_button = WebDriverWait(driver, 10).until(
        EC.element_to_be_clickable((By.XPATH, "//*[@id='add-to-cart-sauce-labs-fleece-jacket']"))
    )
    add_product_button.click()
    time.sleep(2)

    cart_button = WebDriverWait(driver, 10).until(
        EC.element_to_be_clickable((By.XPATH, "//*[@id='shopping_cart_container']/a"))
    )
    cart_button.click()
    time.sleep(2)

    checkout_button = WebDriverWait(driver, 10).until(
        EC.element_to_be_clickable((By.XPATH, "//*[@id='checkout']"))
    )
    checkout_button.click()
    time.sleep(2)

    first_name = WebDriverWait(driver, 10).until(
        EC.presence_of_element_located((By.ID, "first-name"))
    )
    first_name.send_keys("Akash")
    time.sleep(2)

    last_name = WebDriverWait(driver, 10).until(
        EC.presence_of_element_located((By.ID, "last-name"))
    )
    last_name.send_keys("R")
    time.sleep(2)

    postal_code = WebDriverWait(driver, 10).until(
        EC.presence_of_element_located((By.ID, "postal-code"))
    )
    postal_code.send_keys("600049")
    time.sleep(2)

    continue_button = WebDriverWait(driver, 10).until(
        EC.element_to_be_clickable((By.XPATH, "//*[@id='continue']"))
    )
    continue_button.click()
    time.sleep(2)

    try:
        status_message = WebDriverWait(driver, 5).until(
            EC.visibility_of_element_located((By.XPATH, "//*[@id='checkout_summary_container']/div/div[1]/div[3]/div[2]/div[2]/div"))
        )
        if "49.99" in status_message.text:
            finish_button = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.XPATH, "//*[@id='finish']"))
            )
            finish_button.click()
            time.sleep(2)
            thank_you_message = WebDriverWait(driver, 5).until(
                EC.visibility_of_element_located((By.XPATH, "//*[@id='checkout_complete_container']/h2"))
            )
            if "Thank you for your order!" in thank_you_message.text:
                message = "Thank You Checkout Complete page"
            else:
                message = f"Unexpected message: {thank_you_message.text}"
        else:
            message = f"Unexpected product price: {status_message.text}"
    except Exception as e:
        message = f"Error occurred while checking product price: {e}"
    return message
def run_test_cases(driver):
    doc = docx.Document()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    doc_name = f"C:/Users/91951/PycharmProjects/Python selenium/Driver/Screenshots/report_{timestamp}.docx"
    doc.save(doc_name)  # Save the document initially
    test_cases = [
        successful_login,
        invalid_username,
        Order_a_product,
    ]
    for idx, test_case in enumerate(test_cases, start=1):
        message = test_case(driver)
        print(f"Test Case {idx}: {message}")
        pass_screenshot(doc, driver, idx, message, doc_name)
    print(f"Test report saved as {doc_name}")

if __name__ == "__main__":
    driver = initialize_driver()
    try:
        run_test_cases(driver)
    finally:
        driver.quit()
